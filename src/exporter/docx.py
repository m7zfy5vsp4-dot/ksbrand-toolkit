"""Word(docx)导出模块"""

from pathlib import Path
from typing import Optional


class DocxExporter:
    """导出内容为Word文档格式"""

    def export(self, content: str, output_path: str,
               metadata: Optional[dict] = None,
               rag_sources: Optional[list[dict]] = None) -> str:
        """导出为docx文件"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            print("[DocxExporter] python-docx未安装，请执行: pip install python-docx")
            return ""

        filepath = Path(output_path)
        filepath.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft YaHei"
        font.size = Pt(11)

        # 添加元数据
        if metadata:
            for key, value in metadata.items():
                p = doc.add_paragraph()
                run = p.add_run(f"{key}: {value}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)
            doc.add_paragraph("—" * 40)

        # 添加主体内容
        lines = content.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

        # 添加RAG来源标注
        if rag_sources:
            doc.add_paragraph("—" * 40)
            doc.add_heading("RAG来源标注", level=3)
            for i, source in enumerate(rag_sources, 1):
                doc.add_paragraph(
                    f"{i}. 文件: {source.get('source', 'unknown')} | "
                    f"类别: {source.get('category', 'N/A')} | "
                    f"相关度: {source.get('score', 0):.2f}",
                    style="List Number",
                )

        # 添加免责声明
        p = doc.add_paragraph()
        run = p.add_run(
            "【免责声明】以上内容由金山云AI量产助手生成，"
            "仅作内部参考使用，正式发布前请经品牌合规审核。"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(180, 140, 0)

        doc.save(str(filepath))
        return str(filepath)
